import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt

RESULTS_PATHS = [
    os.path.join("reports", "paper_results.csv"),
    os.path.join("reports", "quick_results.csv"),
]
OUT_DOCX = os.path.join("reports", "Battleship_AI_Full_Paper.docx")
FIG_WIN = os.path.join("reports", "fig_win_rate.png")
FIG_MOVES = os.path.join("reports", "fig_avg_moves.png")


def load_results():
    for path in RESULTS_PATHS:
        if os.path.exists(path):
            with open(path, 'r') as f:
                reader = csv.DictReader(f)
                rows = list(reader)
                if rows:
                    return rows, path
    return [], None


def to_float(x, default=0.0):
    try:
        return float(x)
    except Exception:
        return default


def to_int(x, default=0):
    try:
        return int(x)
    except Exception:
        return default


def plot_results(rows):
    if not rows:
        return
    opp = [r['opponent'] for r in rows]
    wins = [to_int(r.get('main_ai_wins', 0)) for r in rows]
    games = [to_int(r.get('games', 0)) or (to_int(r.get('main_ai_wins', 0)) + to_int(r.get('main_ai_losses', 0))) for r in rows]
    win_rates = [ (w/(g if g>0 else 1)) for w,g in zip(wins, games) ]
    avg_moves = [ to_float(r.get('avg_moves_to_win', 0.0)) for r in rows ]

    # Win rate bar
    plt.figure(figsize=(6,4))
    plt.bar(opp, win_rates, color='#2E86C1')
    plt.ylim(0,1)
    plt.ylabel('Win Rate')
    plt.title('AIAgent4 Win Rate by Opponent')
    plt.xticks(rotation=30, ha='right')
    plt.tight_layout()
    plt.savefig(FIG_WIN, dpi=180)
    plt.close()

    # Avg moves bar
    plt.figure(figsize=(6,4))
    plt.bar(opp, avg_moves, color='#7D3C98')
    plt.ylabel('Avg Moves to Win')
    plt.title('AIAgent4 Efficiency by Opponent')
    plt.xticks(rotation=30, ha='right')
    plt.tight_layout()
    plt.savefig(FIG_MOVES, dpi=180)
    plt.close()


def add_title(doc, text):
    title = doc.add_paragraph()
    run = title.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section(doc, heading, paragraphs):
    doc.add_heading(heading, level=1)
    for p in paragraphs:
        para = doc.add_paragraph()
        run = para.add_run(p)
        para.paragraph_format.space_after = Pt(6)


def add_subsection(doc, heading, paragraphs):
    doc.add_heading(heading, level=2)
    for p in paragraphs:
        para = doc.add_paragraph()
        run = para.add_run(p)
        para.paragraph_format.space_after = Pt(4)


def add_results_table(doc, rows):
    if not rows:
        return
    headers = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid'
    hdr = table.rows[0].cells
    for j,h in enumerate(headers):
        hdr[j].text = h
    for r in rows:
        rc = table.add_row().cells
        for j,h in enumerate(headers):
            rc[j].text = str(r.get(h, ""))


def build_paper():
    os.makedirs('reports', exist_ok=True)

    rows, used_path = load_results()
    if rows:
        plot_results(rows)

    doc = Document()

    # Title page
    add_title(doc, "A Probabilistic–Neural–Evolutionary Agent for Battleship: System, Methods, and Evaluation")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Author: Abhinav  |  Date: {datetime.now().strftime('%Y-%m-%d')}")
    run.italic = True

    # Abstract
    add_section(doc, "Abstract", [
        ("We present a complete research platform for the game Battleship and an agent that fuses classical "
         "probabilistic placement counting with a convolutional heatmap prior, Monte Carlo posterior sampling, "
         "information-theoretic scoring, and genetic optimization of blending weights. We document the full "
         "system—including a training pipeline (supervised, optional RL fine-tuning) and a dashboard—then report "
         "benchmarks against heuristic opponents. The approach yields robust win rates and move efficiency. We "
         "release source code, models, and a comprehensive documentation set to facilitate reproducibility.")
    ])

    # Introduction
    add_section(doc, "1. Introduction", [
        ("Battleship is a partially observable search problem in which the agent must infer hidden ship locations "
         "and sequentially query a 10×10 grid. Pure heuristics (e.g., parity search) exploit structure but can stall "
         "without principled posterior reasoning. Our system combines a learned prior with placement enumeration and "
         "Monte Carlo sampling to build a calibrated decision surface, with an information-gain bonus for active search. "
         "Agent 4 further improves the blend via a genetic algorithm (GA)."),
    ])

    # Related Work (concise)
    add_section(doc, "2. Related Work", [
        ("Classical Battleship solvers rely on ship placement counting and parity heuristics (cf. popular analyses such "
         "as DataGenetics). Neural priors (CNNs) learn spatial structure from observations. Monte Carlo methods provide "
         "posterior approximations under constraints. GA optimization is a standard tool for parameter search. Our agent "
         "integrates these strands into a cohesive, reproducible system."),
    ])

    # Problem Formulation
    add_section(doc, "3. Problem Formulation", [
        ("State space consists of the opponent's hidden fleet placement and the agent's observation grid with ternary "
         "values: miss, hit, unknown. The objective is to choose queries (shots) that minimize time-to-sink (equivalently, "
         "maximize win probability subject to move budget). We view the process as Bayesian inference over placements with "
         "actions chosen to maximize expected utility (hit probability or expected information gain)."),
    ])

    # Methods
    add_section(doc, "4. Methods", [
        ("Our policy blends multiple per-cell scorers on the 10×10 board:"),
    ])
    add_subsection(doc, "4.1 Placement Density (D)", [
        ("Enumerate legal placements for the remaining fleet under constraints (misses, sunk-ship cells, and coverage of "
         "active hits). Each valid placement increments the count for its cells. Parity masks accelerate early search by "
         "down-weighting cells that cannot host the smallest remaining ship."),
    ])
    add_subsection(doc, "4.2 Neural Prior (N)", [
        ("A shallow CNN maps the observation tensor (miss, hit, unknown) → per-cell probabilities. The model is trained "
         "supervised on self-play states to predict the true ship mask. Predictions are masked to legal moves and normalized."),
    ])
    add_subsection(doc, "4.3 Monte Carlo Posterior (MC)", [
        ("We sample many full placements consistent with all constraints, optionally increasing samples late in the game. "
         "The normalized occupancy frequency approximates a posterior over ship cells."),
    ])
    add_subsection(doc, "4.4 Information Gain (I)", [
        ("Using the MC occupancy as p(hit), we compute expected entropy reduction per cell. This term promotes shots that "
         "are informative even when immediate hit probability is modest."),
    ])
    add_subsection(doc, "4.5 Target vs Hunt Modes", [
        ("Target: when unsunk hits exist, extend along the inferred ship axis with blocked-end tracking. "
         "Hunt: otherwise, use the blended grid with parity-accelerated density and neural/MC priors."),
    ])
    add_subsection(doc, "4.6 Meta-Weights and GA Optimization", [
        ("We form G = w_D·D̂ + w_N·N̂ + w_MC·MĈ + w_I·Î with a local neighbour boost around unsunk hits. "
         "Agent 4 loads weights from models/ga_weights.json evolved via a GA that maximizes 100×win_rate − avg_moves against a suite of opponents."),
    ])

    # Implementation
    add_section(doc, "5. Implementation", [
        ("Agents: Agent 2 implements D+parity + N + MC with robust target/hunt logic. Agent 3 adds information-gain, "
         "opponent-profiling hooks, and optional graph reasoning; metrics persist to models/. Agent 4 is a thin subclass "
         "of Agent 3 that automatically loads GA-optimized meta-weights."),
        ("Training: generate_dataset.py creates a stream of (state, shipmask) samples; train_heatmap.py trains the CNN; "
         "rl_finetune.py implements a multi-process REINFORCE fine-tune. ga_optimizer.py evolves meta-weights under "
         "memory-safe evaluation."),
        ("Dashboard: battleship_dashboard.py visualizes games, runs batches, and provides analytics (win rates, heatmaps)."),
    ])

    # Experiments
    add_section(doc, "6. Experiments", [
        (f"Results source: {used_path if used_path else 'N/A'}."),
        ("Protocol: AIAgent4 (GA-optimized) vs a small opponent suite (Ultimate, Naive6, AIPlayer2). "
         "We report win rate and average moves to win (lower is better)."),
    ])
    if rows:
        add_results_table(doc, rows)
        if os.path.exists(FIG_WIN):
            doc.add_picture(FIG_WIN, width=Inches(5.8))
        if os.path.exists(FIG_MOVES):
            doc.add_picture(FIG_MOVES, width=Inches(5.8))
    else:
        doc.add_paragraph("No results available; see docs/evaluation.")

    # Limitations & Ablations
    add_section(doc, "7. Limitations and Ablations", [
        ("We anticipate positive contributions from each term (D, N, MC, I) and the target-mode logic. "
         "While ablations are not fully reproduced here, the codebase supports toggling components and reporting deltas. "
         "Limitations include: evaluation primarily versus heuristic baselines; RL fine-tune not yet wired into default agents; "
         "opponent prior integration is partial. See docs/reports/limitations.md."),
    ])

    # Reproducibility
    add_section(doc, "8. Reproducibility", [
        ("Environment: Python 3.9, TensorFlow 2.17.0. See requirements.txt and docs/repro/guide.md for exact commands, seeds, and directories. "
         "We publish models and GA weights under models/."),
    ])

    # Conclusion
    add_section(doc, "9. Conclusion", [
        ("We described a practical, extensible Battleship AI that blends probabilistic structure with learned priors and evolutionary tuning. "
         "Future work: end-to-end opponent prior integration, large-scale ablations with confidence intervals, and comparisons to stronger planners."),
    ])

    doc.save(OUT_DOCX)
    print(f"✅ Full paper written to {OUT_DOCX}")


if __name__ == '__main__':
    build_paper()
