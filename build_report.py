import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Report path
OUT_PATH = os.path.join("reports", "Battleship_AI_Technical_Report.docx")


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr_cells[j].text = h
    for r in rows:
        row_cells = table.add_row().cells
        for j, h in enumerate(headers):
            row_cells[j].text = str(r.get(h, ""))
    table.style = 'Light Grid'
    return table


def read_quick_results(path):
    if not os.path.exists(path):
        return []
    with open(path, 'r') as f:
        reader = csv.DictReader(f)
        return list(reader)


def build_report():
    os.makedirs("reports", exist_ok=True)
    doc = Document()

    # Title
    title = doc.add_paragraph()
    run = title.add_run("On the Probabilistic, Neural, and Evolutionary Design of a Battleship Agent")
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, f"Author: Abhinav (with assistance)", italic=True)
    add_paragraph(doc, f"Date: {datetime.now().strftime('%Y-%m-%d')}")

    add_heading(doc, "Abstract", level=1)
    add_paragraph(doc, (
        "We present a complete research platform for the game of Battleship combining classical "
        "probabilistic placement counting, a convolutional heatmap prior, Monte Carlo posterior blending, "
        "information-theoretic scoring, and genetic optimization of meta-weights (Agent 4). We describe the "
        "system architecture, training pipelines (supervised, RL fine-tuning, GA), and evaluation. A quick "
        "benchmark demonstrates robust performance versus heuristic opponents. We release full documentation "
        "and reproducibility materials."
    ))

    add_heading(doc, "1. Introduction", level=1)
    add_paragraph(doc, (
        "The Battleship problem is a partially observable search task. Our system fuses a neural prior with "
        "probabilistic placement counting and Monte Carlo sampling to infer likely ship locations. Agent 3 "
        "integrates opponent modeling and information gain; Agent 4 loads GA-optimized weights to maximize "
        "win rate while minimizing moves."
    ))

    add_heading(doc, "2. System Overview", level=1)
    add_paragraph(doc, "Environment: 10×10 board, classic ships [5,4,3,3,2]; headless engine for AI-vs-AI.")
    add_paragraph(doc, "Agents: Agent 2 (probabilistic+NN+MC), Agent 3 (adds info gain + opponent modeling + graph), Agent 4 (GA-optimized blend).")
    add_paragraph(doc, "Training: supervised heatmap CNN from self-play data; optional REINFORCE fine-tuning; GA evolution of meta-weights.")

    add_heading(doc, "3. Methods", level=1)
    add_paragraph(doc, "Density counting: enumerate legal placements under misses/sunk constraints; parity reduces early search.")
    add_paragraph(doc, "Neural prior: CNN maps (miss,hit,unknown)→ per-cell ship probability.")
    add_paragraph(doc, "Monte Carlo posterior: sample full-consistent placements to build occupancy frequencies.")
    add_paragraph(doc, "Information gain: expected entropy reduction via MC p(hit).")
    add_paragraph(doc, "Meta-weights: blend grids (density, neural, MC, info, opponent) and boost neighbours of unsunk hits.")
    add_paragraph(doc, "GA: tournament selection, uniform crossover, Gaussian mutation; fitness = 100×win_rate − avg_moves.")

    add_heading(doc, "4. Implementation", level=1)
    add_paragraph(doc, "Codebase covers agents, training scripts, and a Tk dashboard. Logs, datasets, and models reside under data/models/logs.")

    add_heading(doc, "5. Experiments", level=1)
    results = read_quick_results(os.path.join("reports", "quick_results.csv"))
    if results:
        add_paragraph(doc, "Quick benchmark (10 games/opponent; AIAgent4 as main).")
        headers = list(results[0].keys())
        add_table(doc, headers, results)
    else:
        add_paragraph(doc, "No quick results available.")

    add_heading(doc, "6. Discussion", level=1)
    add_paragraph(doc, (
        "Strengths: complete system; robust blend; GA tuning. Limitations: opponent prior not yet fully folded into "
        "inference; RL gains not extensively reported; evaluation vs stronger public baselines is future work."
    ))

    add_heading(doc, "7. Reproducibility", level=1)
    add_paragraph(doc, "See docs/: setup, seeds, commands; scripts for dataset generation, training, and GA.")

    add_heading(doc, "8. Conclusion", level=1)
    add_paragraph(doc, "A practical, extensible Battleship AI stack; future work: end-to-end opponent prior integration, rigorous ablations, larger-scale evaluation.")

    doc.save(OUT_PATH)
    print(f"✅ Report written to {OUT_PATH}")


if __name__ == "__main__":
    build_report()
