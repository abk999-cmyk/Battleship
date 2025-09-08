import os
import csv
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt

# Paths
DOCS_DIR = "docs"
REPORTS_DIR = "reports"
RESULTS_CANDIDATES = [
    os.path.join(REPORTS_DIR, "paper_results.csv"),
    os.path.join(REPORTS_DIR, "quick_results.csv"),
]
OUT_DOCX = os.path.join(REPORTS_DIR, "Battleship_AI_Sheldon_Paper.docx")
FIG_WIN = os.path.join(REPORTS_DIR, "fig_win_rate_full.png")
FIG_MOVES = os.path.join(REPORTS_DIR, "fig_avg_moves_full.png")

# Source docs to stitch
DOC_SOURCES = [
    ("Overview", os.path.join(DOCS_DIR, "overview.md")),
    ("Environment", os.path.join(DOCS_DIR, "environment.md")),
    ("Agent 2", os.path.join(DOCS_DIR, "agents", "agent2.md")),
    ("Agent 3", os.path.join(DOCS_DIR, "agents", "agent3.md")),
    ("Agent 4", os.path.join(DOCS_DIR, "agents", "agent4.md")),
    ("Training Pipelines", os.path.join(DOCS_DIR, "training", "pipelines.md")),
    ("Neural Nets: Heatmap", os.path.join(DOCS_DIR, "nn", "heatmap_model.md")),
    ("Neural Nets: Opponent", os.path.join(DOCS_DIR, "nn", "opponent_model.md")),
    ("Neural Nets: Meta-Learner", os.path.join(DOCS_DIR, "nn", "meta_learner.md")),
    ("Evaluation", os.path.join(DOCS_DIR, "evaluation", "metrics.md")),
    ("Reproducibility", os.path.join(DOCS_DIR, "repro", "guide.md")),
]
APPENDIX_SOURCES = [
    ("Results (Prior Run)", "results.md"),
    ("Methodology", os.path.join(DOCS_DIR, "reports", "methodology.md")),
    ("Data Schemas", os.path.join(DOCS_DIR, "reports", "data_schema.md")),
    ("Logging", os.path.join(DOCS_DIR, "reports", "logging.md")),
    ("Hyperparameters", os.path.join(DOCS_DIR, "reports", "hyperparameters.md")),
    ("Limitations", os.path.join(DOCS_DIR, "reports", "limitations.md")),
]


def read_text(path):
    try:
        with open(path, 'r') as f:
            return f.read()
    except Exception:
        return None


def load_results():
    for p in RESULTS_CANDIDATES:
        if os.path.exists(p):
            with open(p, 'r') as f:
                rows = list(csv.DictReader(f))
            if rows:
                return rows, p
    return [], None


def to_int(x):
    try:
        return int(x)
    except Exception:
        return 0


def to_float(x):
    try:
        return float(x)
    except Exception:
        return 0.0


def ci_wilson(p, n, z=1.96):
    if n == 0:
        return (0.0, 0.0)
    denom = 1 + z**2 / n
    center = (p + z**2/(2*n)) / denom
    margin = (z * ((p*(1-p)/n + z**2/(4*n**2))**0.5)) / denom
    return (max(0.0, center - margin), min(1.0, center + margin))


def plot_results(rows):
    if not rows:
        return
    opp = [r['opponent'] for r in rows]
    wins = [to_int(r.get('main_ai_wins', 0)) for r in rows]
    losses = [to_int(r.get('main_ai_losses', 0)) for r in rows]
    games = [ (to_int(rw.get('games', 0)) or (w + l)) for rw, w, l in zip(rows, wins, losses) ]
    win_rates = [ (w/(g if g>0 else 1)) for w,g in zip(wins, games) ]
    avg_moves = [ to_float(r.get('avg_moves_to_win', 0.0)) for r in rows ]

    # Win rate bar
    plt.figure(figsize=(7,4))
    plt.bar(opp, win_rates, color='#1F77B4')
    plt.ylim(0,1)
    plt.ylabel('Win Rate')
    plt.title('AIAgent4 Win Rate by Opponent (with Wilson 95% CI)')
    # Add error bars
    ci_l = []
    ci_u = []
    for wr, g in zip(win_rates, games):
        l,u = ci_wilson(wr, g)
        ci_l.append(wr - l)
        ci_u.append(u - wr)
    plt.errorbar(opp, win_rates, yerr=[ci_l, ci_u], fmt='none', ecolor='black', capsize=4)
    plt.xticks(rotation=30, ha='right')
    plt.tight_layout()
    plt.savefig(FIG_WIN, dpi=180)
    plt.close()

    # Avg moves bar
    plt.figure(figsize=(7,4))
    plt.bar(opp, avg_moves, color='#9467BD')
    plt.ylabel('Avg Moves to Win')
    plt.title('AIAgent4 Efficiency by Opponent')
    plt.xticks(rotation=30, ha='right')
    plt.tight_layout()
    plt.savefig(FIG_MOVES, dpi=180)
    plt.close()


def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section(doc, heading, text=None):
    doc.add_heading(heading, level=1)
    if text:
        for para in text.split('\n\n'):
            pr = doc.add_paragraph()
            pr.add_run(para)
            pr.paragraph_format.space_after = Pt(6)


def add_subsection(doc, heading, text=None):
    doc.add_heading(heading, level=2)
    if text:
        for para in text.split('\n\n'):
            pr = doc.add_paragraph()
            pr.add_run(para)
            pr.paragraph_format.space_after = Pt(4)


def add_table(doc, rows):
    if not rows:
        return
    headers = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid'
    hdr = table.rows[0].cells
    for j,h in enumerate(headers):
        hdr[j].text = h
    for r in rows:
        rc = table.add_row().cells
        for j,h in enumerate(headers):
            rc[j].text = str(r.get(h, ""))


def main():
    os.makedirs(REPORTS_DIR, exist_ok=True)

    rows, used_path = load_results()
    if rows:
        plot_results(rows)

    doc = Document()

    # Title
    add_title(doc, "A Thorough Exposition of a Probabilistic–Neural–Evolutionary Battleship Agent (Incontrovertibly Superior)")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Author: Abhinav  |  Date: {datetime.now().strftime('%Y-%m-%d')}  |  Data: {os.path.basename(used_path) if used_path else 'results.md'}")
    meta_run.italic = True

    # Abstract (Sheldon-esque precision)
    add_section(doc, "Abstract", (
        "We present a system that optimally (to within reasonable experimental tolerance) navigates Battleship's partial observability by fusing "
        "probabilistic placement densities, a learned convolutional prior, Monte Carlo posterior sampling, and genetic meta-weight optimization. "
        "The resulting agent (hereafter, Agent 4) exhibits substantial superiority over heuristic antagonists, as demonstrated by statistically "
        "defensible win rates and move efficiencies. For the impatient reader: the system works and, dare I say, elegantly."
    ))

    # Stitch core docs
    for title, path in DOC_SOURCES:
        text = read_text(path)
        if text:
            add_section(doc, f"{title}", text)

    # Experiments
    add_section(doc, "Experiments", (
        "We evaluate Agent 4 against an opponent suite (Ultimate, Naive6, Agent2). Each matchup comprises multiple independent games with random ship placements. "
        "Reported metrics include win rate (with Wilson 95% confidence intervals) and average moves-to-win."
    ))
    if rows:
        add_table(doc, rows)
        if os.path.exists(FIG_WIN):
            doc.add_picture(FIG_WIN, width=Inches(6.5))
        if os.path.exists(FIG_MOVES):
            doc.add_picture(FIG_MOVES, width=Inches(6.5))
    else:
        add_subsection(doc, "Results Pending", "No results CSV found; see Appendix for previous large-scale results.")

    # Limitations
    text_lim = read_text(os.path.join(DOCS_DIR, "reports", "limitations.md"))
    add_section(doc, "Limitations", text_lim or "See docs/reports/limitations.md.")

    # Reproducibility
    text_rep = read_text(os.path.join(DOCS_DIR, "repro", "guide.md"))
    add_section(doc, "Reproducibility", text_rep or "See docs/repro/guide.md.")

    # Appendices (prior results and detailed reports)
    doc.add_heading("Appendix", level=1)
    for title, path in APPENDIX_SOURCES:
        text = read_text(path)
        if text:
            add_subsection(doc, title, text)

    doc.save(OUT_DOCX)
    print(f"✅ Sheldon-grade paper written to {OUT_DOCX}")


if __name__ == '__main__':
    main()
